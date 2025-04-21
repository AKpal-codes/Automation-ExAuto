import os
import docx
import pdfplumber
from docx import Document
from huggingface_hub import InferenceClient
from tqdm import tqdm
import time
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from io import BytesIO
from dotenv import load_dotenv
import re

load_dotenv()

client = InferenceClient(
    model="HuggingFaceH4/zephyr-7b-beta",
    token=os.getenv("HUGGINGFACE_TOKEN") 
)

def read_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    if ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    elif ext == ".docx":
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif ext == ".txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
    else:
        raise ValueError("Unsupported file format.")
    return text

def smart_split_text(text, model_max_tokens=3000, buffer_tokens=500):
    words = text.split()
    total_words = len(words)
    max_words_per_chunk = model_max_tokens - buffer_tokens

    if total_words <= max_words_per_chunk:
        return [text] 

    chunks = []
    for i in range(0, total_words, max_words_per_chunk):
        chunk = ' '.join(words[i:i + max_words_per_chunk])
        chunks.append(chunk)

    return chunks


def extract_use_cases(text, max_retries=5):
    prompt_template = f"""
    You are a Business Analyst assistant.

    Analyze the following business-related text and extract detailed, structured Use Cases.

    Document:
    
    {text}

    Format your output exactly like this for each Use Case:
    - Use Case Title:
    - Actor(s):
    - Preconditions:
    - Trigger:
    - Main Flow:
    - Alternative Flows:
    - Postconditions:
    - Notes:

    Do not include any extra text or explanations. Strictly follow the format. Be clean, professional, and concise.
    """

    attempt = 0
    while attempt < max_retries:
        try:
            response = client.text_generation(
                prompt_template,
                max_new_tokens=500,
                temperature=0.7,
                do_sample=True,
                return_full_text=False,
            )
            return response
        except Exception as e:
            print(f"[Retry {attempt+1}/{max_retries}] Error: {e}")
            wait_time = 2 ** attempt  # Exponential backoff
            print(f"Waiting {wait_time} seconds before retrying...")
            time.sleep(wait_time)
            attempt += 1

    raise RuntimeError("Failed to extract use cases after multiple retries.")

def is_valid_email(email):
    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(email_regex, email)

def read_recipients_from_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            recipient_emails = [line.strip() for line in file if line.strip()]
        valid_emails = [email for email in recipient_emails if is_valid_email(email)]
        invalid_emails = set(recipient_emails) - set(valid_emails)
        if invalid_emails:
            print(f"Warning: The following emails are invalid and will be ignored: {', '.join(invalid_emails)}")
        return valid_emails
    except Exception as e:
        print(f"Error reading recipients file: {e}")
        return []

SMTP_SERVER = os.getenv("SMTP_SERVER")
SMTP_PORT = int(os.getenv("SMTP_PORT"))

def send_docx_via_email(structured_texts, recipient_emails, sender_email, sender_password):
    # Create the document in memory
    doc = Document()
    doc.add_heading('Extracted Use Cases', level=0)

    for idx, structured_text in enumerate(structured_texts, start=1):
        doc.add_heading(f'Chunk {idx}', level=1)
        for block in structured_text.split('\n'):
            if block.strip() == "":
                continue
            if block.startswith("- Use Case Title"):
                doc.add_heading(block.replace("- ", ""), level=2)
            elif block.startswith("- "):
                para = doc.add_paragraph(block)
                para.style = 'List Bullet'
            else:
                doc.add_paragraph(block)

    # Save the document to a BytesIO object
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # Email setup
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = ", ".join(recipient_emails) 
    msg['Subject'] = "Extracted Use Cases Document"

    body = "Please find the extracted use cases document attached."
    msg.attach(MIMEText(body, 'plain'))

    # Attach the document
    attachment = MIMEBase('application', 'octet-stream')
    attachment.set_payload(doc_stream.read())
    encoders.encode_base64(attachment)
    attachment.add_header('Content-Disposition', 'attachment', filename="Extracted_Use_Cases.docx")
    msg.attach(attachment)

    # Send the email
    try:
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
        print(f"Email sent successfully to {', '.join(recipient_emails)}")
    except Exception as e:
        print(f"Failed to send email: {e}")

def main(input_path, recipient_path):
    print("Reading document...")
    raw_text = read_document(input_path)
    print("Splitting document into chunks...")
    chunks = smart_split_text(raw_text)
    print(f"Total Chunks: {len(chunks)}")

    all_structured_outputs = []

    for chunk in tqdm(chunks, desc="Processing chunks"):
        structured_text = extract_use_cases(chunk)
        all_structured_outputs.append(structured_text)

    print("Reading recipients...")
    recipient_emails = read_recipients_from_file(recipient_path)
    print("Sending structured output via email...")
    username = os.getenv("EMAIL_USERNAME")
    password = os.getenv("EMAIL_PASSWORD")
    send_docx_via_email(
        structured_texts=all_structured_outputs,
        recipient_emails=recipient_emails,
        sender_email=username,
        sender_password=password
    )
    print("Done! Structured document sent to the recipients.")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Extract structured use cases from a business document.")
    parser.add_argument("--input", required=True, help="Path to input document (PDF, DOCX, or TXT)")
    parser.add_argument("--recipients", required=True, help="Path to recipients mails file")
    args = parser.parse_args()

    main(args.input, args.recipients)